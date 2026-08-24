import csv, json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

# Read latest CSV
csv_path = '/Users/xiaozuo/Downloads/tcl_price_reports/tcl_prices_2026-08-24T16-44-29.csv'
rows = []
with open(csv_path, newline='', encoding='utf-8') as f:
    reader = csv.DictReader(f)
    for row in reader:
        rows.append(row)

# Filter Amazon only
amazon_rows = [r for r in rows if r['Platform'] == 'Amazon']

doc = Document()

# Header
title = doc.add_heading('TCL 电视价格监控报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('报告时间: 2026年8月24日 16:44')
doc.add_paragraph('数据来源: Amazon.com 共20个SKU，4个Best Buy因封锁未抓取')
doc.add_paragraph('')
doc.add_paragraph('注: Best Buy 4个SKU因HTTP2封锁未能抓取，详见附录。')
doc.add_paragraph('')

# Table
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
headers = ['序号', '产品型号', 'DCT定价', 'Amazon实价', '折扣', '状态']
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    for p in hdr_cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

for idx, r in enumerate(amazon_rows, 1):
    row = table.add_row()
    row.cells[0].text = str(idx)
    row.cells[1].text = r['Product'].strip()
    row.cells[2].text = r['DCT_Pricing']

    # Price
    price_val = r['Price_Value']
    if price_val:
        try:
            pv = float(price_val)
            row.cells[3].text = '${:,.2f}'.format(pv)
        except:
            row.cells[3].text = r['Price_Text'][:30] if r['Price_Text'] else 'N/A'
    else:
        row.cells[3].text = 'N/A'

    # Discount
    dct = r['DCT_Pricing']
    pv = r['Price_Value']
    if dct and pv:
        try:
            dct_f = float(dct.replace('$', '').replace(',', ''))
            pv_f = float(pv)
            if dct_f > 0:
                discount = (dct_f - pv_f) / dct_f * 100
                row.cells[4].text = '{:+.1f}%'.format(discount)
            else:
                row.cells[4].text = 'N/A'
        except:
            row.cells[4].text = 'N/A'
    else:
        row.cells[4].text = 'N/A'

    row.cells[5].text = r['Status']

doc.add_paragraph('')

# Best Buy section
doc.add_heading('Best Buy 未抓取SKU (HTTP2封锁)', level=2)
bb_rows = [r for r in rows if r['Platform'] == 'Best Buy']
if bb_rows:
    bb_table = doc.add_table(rows=1, cols=4)
    bb_table.style = 'Table Grid'
    bb_hdr = bb_table.rows[0].cells
    for i, h in enumerate(['产品型号', 'URL', 'DCT定价', '备注']):
        bb_hdr[i].text = h
        for p in bb_hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r in bb_rows:
        row = bb_table.add_row()
        row.cells[0].text = r['Product']
        row.cells[1].text = (r['URL'][:60] + '...') if len(r['URL']) > 60 else r['URL']
        row.cells[2].text = r['DCT_Pricing'] if r['DCT_Pricing'] else 'N/A'
        row.cells[3].text = 'Best Buy HTTP2封锁 — 需BESTBUY_API_KEY或搜索页面抓取'

# Footer
doc.add_paragraph('')
p = doc.add_paragraph('— 报告结束 —')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
out_path = '/Users/xiaozuo/Downloads/TCL_价格监控报告_2026-08-24.docx'
doc.save(out_path)
print('Saved: ' + out_path)
print('Size: {} bytes'.format(os.path.getsize(out_path)))
