#!/usr/bin/env python3
"""
Template-based resume and cover letter .docx generator.
Uses an existing .docx as template, replaces content while preserving formatting.

Usage:
  python3 generate-resume.py --type resume --template <path> --content '<json>' --output <path>
  python3 generate-resume.py --type cover_letter --template <path> --content '<json>' --output <path>
"""

import argparse
import json
import sys
import os
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def generate_resume(template_path, content_json, output_path):
    """Generate a tailored resume from template + JSON content."""
    content = json.loads(content_json) if isinstance(content_json, str) else content_json

    # Create new doc with same page setup as template
    template = Document(template_path)
    doc = Document()

    # Copy page setup from template
    template_section = template.sections[0]
    section = doc.sections[0]
    section.top_margin = template_section.top_margin
    section.bottom_margin = template_section.bottom_margin
    section.left_margin = template_section.left_margin
    section.right_margin = template_section.right_margin
    section.page_width = template_section.page_width
    section.page_height = template_section.page_height

    # Copy styles from template if possible
    try:
        for style in template.styles:
            pass  # Styles are read-only, we'll use matching style names
    except Exception:
        pass

    # --- Build Resume Content ---

    # 1. Name header
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(content.get('name', 'BARRON ZUO'))
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_run.font.name = 'Calibri'
    name_para.space_after = Pt(2)

    # 2. Contact line
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run(content.get('contact', 'San Francisco, CA | +1 909-413-2840 | xz429@cornell.edu'))
    contact_run.font.size = Pt(9)
    contact_run.font.name = 'Calibri'
    contact_para.space_after = Pt(6)

    # 3. Executive Summary
    doc.add_heading('EXECUTIVE SUMMARY', level=1)
    summary_para = doc.add_paragraph(content.get('executive_summary', ''))
    summary_para.style.font.size = Pt(10)
    for run in summary_para.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    # 4. Core Competencies (table) — omitted entirely if not provided
    competencies = content.get('competencies', [])
    if competencies:
        doc.add_heading('CORE COMPETENCIES', level=1)
        table = doc.add_table(rows=len(competencies), cols=2)
        table.autofit = True
        for i, row_data in enumerate(competencies):
            for j, cell_text in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = cell_text
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9.5)
                        run.font.name = 'Calibri'

    # 5. Professional Experience
    doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
    for exp in content.get('experience', []):
        # Company line
        company_para = doc.add_paragraph()
        company_run = company_para.add_run(exp.get('company', ''))
        company_run.bold = True
        company_run.font.size = Pt(10.5)
        company_run.font.name = 'Calibri'
        company_para.space_after = Pt(0)

        # Role line
        role_para = doc.add_paragraph()
        role_run = role_para.add_run(exp.get('role', ''))
        role_run.italic = True
        role_run.font.size = Pt(10)
        role_run.font.name = 'Calibri'
        role_para.space_after = Pt(2)

        # Bullet points
        for bullet in exp.get('bullets', []):
            bullet_para = doc.add_paragraph(bullet, style='List Bullet')
            for run in bullet_para.runs:
                run.font.size = Pt(9.5)
                run.font.name = 'Calibri'
            bullet_para.space_after = Pt(1)

    # 6. Education
    doc.add_heading('EDUCATION', level=1)
    for edu in content.get('education', []):
        edu_para = doc.add_paragraph(edu, style='List Bullet')
        for run in edu_para.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

    # Save
    out_dir = os.path.dirname(os.path.abspath(output_path))
    os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)
    print(json.dumps({'success': True, 'output': output_path, 'message': 'Resume generated'}))


def generate_cover_letter(template_path, content_json, output_path):
    """Generate a tailored cover letter from template + JSON content."""
    content = json.loads(content_json) if isinstance(content_json, str) else content_json

    # Create new doc with same page setup as template
    template = Document(template_path)
    doc = Document()

    # Copy page setup
    template_section = template.sections[0]
    section = doc.sections[0]
    section.top_margin = template_section.top_margin
    section.bottom_margin = template_section.bottom_margin
    section.left_margin = template_section.left_margin
    section.right_margin = template_section.right_margin

    # --- Build Cover Letter Content ---

    # 1. Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_lines = content.get('header', 'BARRON ZUO\nSan Francisco, CA | +1 909-413-2840 | xz429@cornell.edu').split('\n')
    for i, line in enumerate(header_lines):
        run = header_para.add_run(line)
        if i == 0:
            run.bold = True
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(9)
        run.font.name = 'Calibri'
        if i < len(header_lines) - 1:
            header_para.add_run('\n')

    # 2. Date
    doc.add_paragraph('')
    date_para = doc.add_paragraph(content.get('date', ''))
    for run in date_para.runs:
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'

    # 3. Recipient
    doc.add_paragraph('')
    recipient_lines = content.get('recipient', '').split('\n')
    for line in recipient_lines:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.name = 'Calibri'

    # 4. Salutation
    doc.add_paragraph('')
    sal_para = doc.add_paragraph(content.get('salutation', 'Dear Hiring Team,'))
    for run in sal_para.runs:
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'

    # 5. Body paragraphs
    for para_text in content.get('paragraphs', []):
        p = doc.add_paragraph(para_text)
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.name = 'Calibri'
        p.space_after = Pt(6)

    # 6. Sign off
    doc.add_paragraph('')
    sign_off_lines = content.get('sign_off', 'Sincerely,\n\nBarron Zuo').split('\n')
    for line in sign_off_lines:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.name = 'Calibri'

    # Save
    out_dir = os.path.dirname(os.path.abspath(output_path))
    os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)
    print(json.dumps({'success': True, 'output': output_path, 'message': 'Cover letter generated'}))



# Post-2018 tools/products that must never appear attached to GSV/UGL (2010-2016)
# or WeWork (2019-2020) bullets. Not exhaustive — a denylist is a backstop, not a
# substitute for the anachronism-check rule in resume-prompt.md / cover-letter-prompt.md —
# but it catches the exact failure class that already shipped once (Issue 18: a
# "Clay-style enrichment" claim on a 2010-2016 bullet; Clay did not exist until years later).
ANACHRONISTIC_TOOLS = [
    'clay', 'chatgpt', 'notion ai', 'gpt-4', 'gpt-5', 'claude ai', 'copilot',
    'midjourney', 'gemini ai', 'perplexity ai', 'clearbit', 'apollo.io',
    'hubspot ai', 'salesforce einstein', 'segment.io cdp', 'rockerbox',
]
PRE_2018_SECTION_MARKERS = ['gsv global tech', 'ugl consulting', 'wework labs', 'wework']


def _check_anachronisms(paragraphs):
    """Flag (not silently fix) any post-2018 tool name appearing inside the
    pre-2018 EXPERIENCE SECTION (GSV/UGL 2010-2016, WeWork 2019-2020) — i.e. in
    the bullets that follow a short section-header paragraph naming that company,
    up until the next company header. Deliberately does NOT scan the Executive
    Summary or any other paragraph that merely *mentions* GSV/WeWork in passing
    while narrating present-day experience (that produced false positives — a
    summary sentence like "...backed by consulting at GSV Global Tech..." sitting
    in the same paragraph as an unrelated modern-tool reference elsewhere in that
    same long summary paragraph is not an anachronism)."""
    problems = []
    in_pre2018_section = False
    for para in paragraphs:
        stripped = para.strip()
        lower = stripped.lower()
        is_header_like = len(stripped) < 90  # company/role header lines are short; prose bullets/summary are not
        if is_header_like and any(marker in lower for marker in PRE_2018_SECTION_MARKERS):
            in_pre2018_section = True
            continue
        if is_header_like and in_pre2018_section:
            # a short line that ISN'T a pre-2018 marker while we're inside the
            # section is very likely the NEXT company's header — section ends
            in_pre2018_section = False
        if not in_pre2018_section:
            continue
        for tool in ANACHRONISTIC_TOOLS:
            if tool in lower:
                problems.append(
                    f'possible anachronism: "{tool}" appears in a bullet under a pre-2018 '
                    f'section header (GSV 2010-2016 / WeWork 2019-2020) — verify this tool '
                    f'existed and was in use at that time, or replace with generic process '
                    f'language before shipping. Bullet: "{stripped[:150]}"'
                )
    return problems


def validate_output(doc_type, output_path):
    """Hard content gate: abort (delete output, exit non-zero) if the generated
    document is too thin to be a real resume/cover letter. Catches cases where
    the caller passed near-empty JSON (e.g. an empty paragraphs/bullets list)
    and the generator silently produced a near-blank .docx."""
    doc = Document(output_path)
    text = '\n'.join(p.text for p in doc.paragraphs)
    problems = []

    if doc_type == 'cover_letter':
        if len(text) < 1200:
            problems.append(f'cover letter body is only {len(text)} chars (minimum 1200)')
        # body paragraphs = everything between the salutation and the sign-off
        lower = text.lower()
        sal_idx = lower.find('dear ')
        sign_idx = lower.rfind('sincerely')
        if sal_idx != -1 and sign_idx != -1 and sign_idx > sal_idx:
            body = text[sal_idx:sign_idx]
            non_empty_paras_in_body = [
                p.text.strip() for p in doc.paragraphs
                if p.text.strip() and lower.find(p.text.strip().lower()) >= sal_idx
                and lower.find(p.text.strip().lower()) < sign_idx
            ]
            # require at least 3 real body paragraphs with meaningful length each —
            # catches the exact Issue 11 failure (scaffold-only letter: greeting
            # immediately followed by sign-off, near-zero body content) that a
            # pure char-count/byte-size check missed because docx boilerplate
            # dominates file size regardless of body content.
            substantive = [p for p in non_empty_paras_in_body if len(p) > 100]
            if len(substantive) < 3:
                problems.append(
                    f'only {len(substantive)} substantive body paragraphs (>100 chars each) '
                    f'found between salutation and sign-off (minimum 3) — likely a blank/scaffold '
                    f'letter with greeting immediately followed by sign-off'
                )
            if body.count('\n') < 3 and len(body) < 800:
                problems.append('fewer than 3 body paragraphs between salutation and sign-off')
            # Issue 10: LLM wrote its own sign-off inside the body, duplicating the
            # template's separate sign_off field — appears twice in the rendered doc.
            if body.count('sincerely') >= 1:
                problems.append(
                    'a "Sincerely" sign-off was found INSIDE the body paragraphs — the body '
                    'must end at the call-to-action sentence; the template appends its own '
                    'sign_off field separately, so a body sign-off causes a duplicate'
                )
        else:
            problems.append('could not locate both a salutation ("Dear ...") and a sign-off ("Sincerely")')
        problems.extend(_check_anachronisms([p.text for p in doc.paragraphs]))
    elif doc_type == 'resume':
        if len(text) < 2000:
            problems.append(f'resume body is only {len(text)} chars (minimum 2000)')
        required_companies = ['alibaba', 'next2market', 'wework', 'indiegogo', 'gsv']
        lower = text.lower()
        missing = [c for c in required_companies if c not in lower]
        if missing:
            problems.append(f'missing required companies in body: {missing}')
        problems.extend(_check_anachronisms([p.text for p in doc.paragraphs]))

    return problems


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Generate resume/cover letter .docx from template + JSON')
    parser.add_argument('--type', required=True, choices=['resume', 'cover_letter'], help='Document type')
    parser.add_argument('--template', required=True, help='Path to template .docx file')
    parser.add_argument('--content', required=True, help='JSON string with content')
    parser.add_argument('--output', required=True, help='Output .docx file path')
    args = parser.parse_args()

    try:
        if args.type == 'resume':
            generate_resume(args.template, args.content, args.output)
        else:
            generate_cover_letter(args.template, args.content, args.output)

        problems = validate_output(args.type, args.output)
        if problems:
            os.remove(args.output)
            print(json.dumps({
                'success': False,
                'error': 'Generated document failed content validation and was deleted — do not upload it.',
                'problems': problems,
            }), file=sys.stderr)
            sys.exit(1)
    except Exception as e:
        print(json.dumps({'success': False, 'error': str(e)}), file=sys.stderr)
        sys.exit(1)
        sys.exit(1)
